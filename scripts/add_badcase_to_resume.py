# -*- coding: utf-8 -*-
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path(r"C:\Users\h2629\Downloads\胡雄java开发简历.docx")
OUT = Path(r"C:\Users\h2629\Downloads\胡雄java开发简历_新增BadCaseDoctor项目.docx")


NEW_PROJECT_PARAS = [
    "2026.02-至今 BadCase Doctor AI质量治理与智能测试平台 个人项目",
    "项目概述：BadCase Doctor 是一个面向 AI 应用质量治理和智能测试的工作台产品，主要解决大模型应用上线后 BadCase 难发现、难归因、难闭环的问题。系统支持项目管理、Bug 管理、BadCase 管理、测试用例管理、AI 对话式操作、自动测试、浏览器执行、终端执行、指标采集和质量分析，目标是把 AI 应用质量问题沉淀为可追踪、可复盘、可持续优化的质量资产。",
    "核心能力：支持由测试用例驱动自动测试，也支持用户通过口头描述测试目标，由 Agent 自动拆解测试步骤，并借助 browser-use / Playwright 类浏览器自动化能力执行点击、输入、跳转、截图、观察和结果判断。测试失败后可自动生成候选 Bug 或 BadCase，人工确认后进入项目闭环。",
    "技术栈：Java、Spring Boot、Spring AI、Micrometer、Prometheus、Python Flask、MySQL、Redis、Elasticsearch、Electron、Vue3、Go 本地代理、SSE 流式协议、browser-use / Playwright 类浏览器自动化。",
    "负责的工作：",
    "1. 独立完成产品架构设计和核心功能开发，设计 Bug、BadCase、测试用例、项目、成员、评论、卡片等核心业务模型，并形成统一项目协作闭环。",
    "2. 设计并实现 AI Agent 工具调用流程，支持自然语言查询、创建、修改、指派和总结问题，将传统表单操作升级为对话式操作。",
    "3. 设计自动测试能力，支持测试用例驱动和口头描述驱动两种方式，通过浏览器自动执行流程、采集截图证据、生成执行结果，并将失败结果转为候选 Bug / BadCase。",
    "4. 设计 BadCase 质量归因方案，从模型参数、输入质量、知识库召回、工具调用、服务稳定性、输出质量等维度分析 AI 问题来源。",
    "5. 设计 Java / Python SDK 与 Prometheus 指标采集方案，规划请求耗时、错误率、流式输出、工具调用、BadCase 类型等指标，为后续 AI 质量运维和看板分析提供数据基础。",
    "6. 实现桌面端工作台与本地代理能力，支持终端执行、项目上下文加载、SSE 流式展示、沙箱预览和多工具协作，提高测试、研发、运维处理问题的效率。",
    "项目亮点：将缺陷管理、AI BadCase 治理、自动测试、浏览器执行、指标观测和智能研发规划整合到一个工作台中，具备商业化 SaaS / 私有化部署方向，可用于企业 AI 应用上线后的质量治理和持续运营。",
    "",
]


def copy_para_format(src, dst):
    dst.style = src.style
    dst.paragraph_format.alignment = src.paragraph_format.alignment
    dst.paragraph_format.left_indent = src.paragraph_format.left_indent
    dst.paragraph_format.right_indent = src.paragraph_format.right_indent
    dst.paragraph_format.first_line_indent = src.paragraph_format.first_line_indent
    dst.paragraph_format.space_before = src.paragraph_format.space_before
    dst.paragraph_format.space_after = src.paragraph_format.space_after
    dst.paragraph_format.line_spacing = src.paragraph_format.line_spacing


def apply_font(paragraph, bold=False):
    for run in paragraph.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
        run.font.bold = bold


def insert_before(anchor, text, template):
    p = anchor.insert_paragraph_before(text)
    copy_para_format(template, p)
    apply_font(p, bold=text.startswith("2026.02"))
    return p


def main():
    doc = Document(str(SRC))
    cell = doc.tables[0].cell(2, 0)

    anchor = None
    for p in cell.paragraphs:
        if "2024-8-至今 相闻AI交通问答" in p.text:
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("未找到第一个项目经历锚点：2024-8-至今 相闻AI交通问答")

    template = anchor
    for text in NEW_PROJECT_PARAS:
        insert_before(anchor, text, template)

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
