# -*- coding: utf-8 -*-
"""完善「相闻 AI 交通问答」项目经历并导出 PDF。"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SRC = Path(r"C:\Users\h2629\Desktop\胡雄java及Agent开发简历.docx")
OUT_DOCX = SRC
OUT_PDF = Path(r"C:\Users\h2629\Desktop\胡雄java及Agent开发简历_更新.pdf")
OUT_PDF_MAIN = Path(r"C:\Users\h2629\Desktop\胡雄java及Agent开发简历.pdf")

TITLE = "2024.8-2026.6 相闻 AI 交通问答 百度的产品已上线"
OVERVIEW = (
    "相闻 AI 交通问答是一款面向交通出行场景的大模型问答 SaaS 系统（百度侧产品已上线），"
    "已落地西安永安君、石家庄石头君等客户。系统覆盖知识库、知识平台、知识底座、会话管理、"
    "AI 意图识别与工具调用等模块，支持基于企业文档的智能问答与业务办理辅助。"
)
TECH = (
    "使用的技术栈: springcloud, springboot, mysql, elasticsearch, "
    "千帆大模型, boss 文件系统, python langchain, spring ai, micrometer"
)
DUTIES = [
    "1. 知识库相关功能开发与重构，优化召回策略与相关度排序，将文档召回精确度从 90% 提升到 95%。",
    "2. 整体接口性能优化，将主要问答接口耗时从 1s 以上优化到约 300ms。",
    "3. 动态提示词优化：结合上下文与召回结果构建 System Prompt / Few-shot，提升问答准确度。",
    "4. 设计并实现 Agent 工具调用链路追踪（基于 Micrometer Observation），观测工具调用耗时与成功率，支撑 Agent 效果评测。",
    "5. 参与会话、AI 意图等模块开发，保障多轮对话与意图路由稳定可用。",
]


def clear_runs(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)


def set_text(paragraph, text, bold=None):
    clear_runs(paragraph)
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold
    return paragraph


def insert_after(paragraph, template, text):
    new_p = deepcopy(template._p)
    for child in list(new_p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph):
    p = paragraph._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def main():
    doc = Document(str(SRC))

    title_i = None
    for i, p in enumerate(doc.paragraphs):
        if "相闻" in p.text and "交通问答" in p.text:
            title_i = i
            break
    if title_i is None:
        raise RuntimeError("未找到相闻项目标题")

    # 当前结构：title, overview, tech, duty1, duty2, duty3, (empty...), 五育
    set_text(doc.paragraphs[title_i], TITLE, bold=True)
    set_text(doc.paragraphs[title_i + 1], OVERVIEW, bold=False)
    set_text(doc.paragraphs[title_i + 2], TECH, bold=False)

    # 收集标题后、五育前的职责/空段落，统一替换
    end_i = None
    for i in range(title_i + 1, len(doc.paragraphs)):
        if "五育" in doc.paragraphs[i].text:
            end_i = i
            break
    if end_i is None:
        raise RuntimeError("未找到五育项目作为边界")

    # 模板：用原第 1 条职责段落版式
    item_tpl = doc.paragraphs[title_i + 3]

    # 删除 title+3 到 五育 之前的所有段落（从后往前删，避免索引错乱）
    to_delete = list(range(title_i + 3, end_i))
    # 重新取 paragraph 对象（删除时用当前列表快照）
    paras_snapshot = list(doc.paragraphs)
    for i in reversed(to_delete):
        delete_paragraph(paras_snapshot[i])

    # 在 tech 段落后插入职责
    # 删除后重新定位
    title_i = None
    for i, p in enumerate(doc.paragraphs):
        if "相闻" in p.text and "交通问答" in p.text:
            title_i = i
            break
    tech_p = doc.paragraphs[title_i + 2]
    # item_tpl 可能已被删；找任意仍存在的 1. 开头段落作模板，否则用 tech
    item_tpl = tech_p
    for p in doc.paragraphs:
        if p.text.strip().startswith("1.") and ("规则引擎" in p.text or "设备心跳" in p.text or "知识库" in p.text):
            item_tpl = p
            break
    # 优先用 esl/beacon 条目缩进
    for p in doc.paragraphs:
        if p.text.strip().startswith("1.规则引擎"):
            item_tpl = p
            break

    last = tech_p
    # insert_after 每次插到 last 后面；要保持顺序，依次插在上一条后
    for text in DUTIES:
        last = insert_after(last, item_tpl, text)

    doc.save(str(OUT_DOCX))
    print("DOCX:", OUT_DOCX)

    # 校验
    for p in doc.paragraphs:
        if "相闻" in p.text or p.text.strip()[:2] in ("1.", "2.", "3.", "4.", "5.") and "知识库" in p.text or "Micrometer" in p.text or "提示词" in p.text:
            if "相闻" in p.text or "知识库" in p.text or "Micrometer" in p.text or "提示词" in p.text or "接口性能" in p.text or "会话" in p.text:
                print(">", p.text[:90])

    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        for out in (OUT_PDF, OUT_PDF_MAIN):
            try:
                if out.exists():
                    out.unlink()
            except PermissionError:
                print("skip locked", out)
                continue
            d = word.Documents.Open(str(OUT_DOCX), ReadOnly=True)
            d.SaveAs(str(out), FileFormat=17)
            d.Close(False)
            print("PDF:", out, out.stat().st_size)
    finally:
        word.Quit()


if __name__ == "__main__":
    main()
