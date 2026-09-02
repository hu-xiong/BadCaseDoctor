# -*- coding: utf-8 -*-
"""基于「胡雄java及Agent开发简历」完善东亚银行项目经历并导出 PDF。"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

SRC = Path(r"C:\Users\h2629\Downloads\胡雄java及Agent开发简历.docx")
OUT_DOCX = Path(r"C:\Users\h2629\Desktop\胡雄java及Agent开发简历.docx")
OUT_PDF = Path(r"C:\Users\h2629\Desktop\胡雄java及Agent开发简历.pdf")

OVERVIEW = (
    "东亚银行手机银行与官网支撑系统，覆盖转账汇款、支付、金融理财、会员管理、"
    "财务审批与资金流转等核心业务。含会员积分与「万里行」类权益体系（消费/活动累计积分、"
    "积分兑换、权益发放与核销），支持香港地区青年派发港币等活动能力，并落地 App 端二维码"
    "电子支付，保障 Web 与 App 业务协同运转。"
)
TECH = "技术栈: ionic angular, springboot, oracle"
DUTY_HEADER = "负责的工作："
DUTIES = [
    "1.负责 Web 端业务模块日常开发与维护，跟进转账、支付、会员、理财等接口与页面联调，处理线上问题与需求迭代。",
    "2.参与会员积分 / 「万里行」类权益模块：积分规则配置、消费/活动累计、积分查询与流水、兑换与核销、权益到账状态同步，保证积分增减可追溯。",
    "3.参与 App 手机银行香港地区二维码电子支付：支付下单、扫码/被扫、结果回调与订单状态同步，保障支付闭环。",
    "4.参与财务审批与支付流转：按审批节点推进单据状态，保证审批流与资金操作可审计。",
    "5.基于 Oracle 完成积分流水、订单、会员等数据的查询与持久化，处理事务一致性与常见慢查询。",
    "6.配合测试与业务方联调验收，沉淀接口与问题记录，保障版本按期上线。",
]


def clear_runs(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)


def set_text(paragraph, text, bold=None):
    clear_runs(paragraph)
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold
    return paragraph


def copy_paragraph_format(src, dst):
    spf, dpf = src.paragraph_format, dst.paragraph_format
    dpf.alignment = spf.alignment
    dpf.left_indent = spf.left_indent
    dpf.right_indent = spf.right_indent
    dpf.first_line_indent = spf.first_line_indent
    dpf.space_before = spf.space_before
    dpf.space_after = spf.space_after
    dpf.line_spacing = spf.line_spacing
    try:
        dst.style = src.style
    except Exception:
        pass


def insert_after(paragraph, template, text):
    new_p = deepcopy(template._p)
    for child in list(new_p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def find_para(doc, predicate):
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text):
            return i, p
    raise RuntimeError("未找到目标段落")


def update_docx():
    doc = Document(str(SRC))

    title_i, title_p = find_para(
        doc, lambda t: "东亚银行手机银行" in t and "官网" not in t and len(t.strip()) < 40
    )
    overview_p = doc.paragraphs[title_i + 1]
    tech_p = doc.paragraphs[title_i + 2]
    duty_p = doc.paragraphs[title_i + 3]

    # 标题格式对齐：参考 beaconcloud 标题缩进风格（短标题略大缩进可保留加粗）
    set_text(title_p, "东亚银行手机银行", bold=True)

    # 概述 / 技术栈：沿用原段落版式，只替换文字
    set_text(overview_p, OVERVIEW, bold=False)
    set_text(tech_p, TECH, bold=False)

    # 「负责的工作」对齐 beaconcloud / esl 段落格式
    duty_tpl = None
    item_tpl = None
    for p in doc.paragraphs:
        if p.text.strip() in ("负责的工作：", "负责的工作:"):
            duty_tpl = p
        if p.text.strip().startswith("1.") and "规则引擎" in p.text:
            item_tpl = p
        if p.text.strip().startswith("1.") and "设备心跳" in p.text and item_tpl is None:
            item_tpl = p
    if duty_tpl is None or item_tpl is None:
        # fallback：用当前 duty 段
        duty_tpl = duty_p
        item_tpl = duty_p

    set_text(duty_p, DUTY_HEADER, bold=False)
    copy_paragraph_format(duty_tpl, duty_p)

    # 删除东亚银行段落后多余旧内容（若有）
    # 在 duty 后插入职责条目
    last = duty_p
    for text in DUTIES:
        last = insert_after(last, item_tpl, text)

    doc.save(str(OUT_DOCX))
    print("DOCX:", OUT_DOCX)
    return OUT_DOCX


def export_pdf(docx_path: Path, pdf_path: Path):
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        # 若目标 PDF 已打开可能导致失败，先写临时再替换
        tmp = pdf_path.with_suffix(".tmp.pdf")
        if tmp.exists():
            tmp.unlink()
        doc = word.Documents.Open(str(docx_path))
        doc.SaveAs(str(tmp), FileFormat=17)
        doc.Close(False)
        if pdf_path.exists():
            pdf_path.unlink()
        tmp.rename(pdf_path)
    finally:
        word.Quit()
    print("PDF:", pdf_path, "size", pdf_path.stat().st_size)


def main():
    docx_path = update_docx()
    export_pdf(docx_path, OUT_PDF)


if __name__ == "__main__":
    main()
