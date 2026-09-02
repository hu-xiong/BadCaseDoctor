# -*- coding: utf-8 -*-
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
import win32com.client

SRC = Path(r"C:\Users\h2629\Desktop\胡雄java及Agent开发简历.docx")
OUT_PDF = Path(r"C:\Users\h2629\Desktop\胡雄java及Agent开发简历_更新.pdf")
OUT_PDF_MAIN = Path(r"C:\Users\h2629\Desktop\胡雄java及Agent开发简历.pdf")
DL = Path(r"C:\Users\h2629\Downloads\胡雄java及Agent开发简历.docx")

TITLE = "2024.8-2026.6 相闻 AI 交通问答 百度的产品已上线"
OVERVIEW = (
    "相闻 AI 交通问答是一款面向交通出行场景的大模型问答 SaaS 系统（百度侧产品已上线），"
    "已落地西安永安君、石家庄石头君等客户。系统覆盖知识库（文档解析/分片入库）、知识平台、知识底座、"
    "会话管理、AI 意图识别、工具调用与拒答策略等模块，支持基于企业文档的智能问答与业务办理辅助。"
)
TECH = (
    "使用的技术栈: springcloud, springboot, mysql, elasticsearch, "
    "千帆大模型, boss 文件系统, python langchain, spring ai, micrometer"
)
DUTIES = [
    "1. 知识库开发与重构：落地文档解析、分片入库与元数据管理，支撑大规模交通领域语料接入。",
    "2. 优化召回链路，落地关键词 + 向量混合检索与相关度排序，将文档召回精确度从 90% 提升到 95%。",
    "3. 整体接口性能优化，将主要问答接口耗时从 1s 以上优化到约 300ms。",
    "4. 动态提示词优化：结合上下文与召回结果构建 System Prompt / Few-shot，提升问答准确度。",
    "5. 设计拒答策略：对低置信召回、越权/超范围问题进行拒答或澄清引导，降低幻觉与错误答复风险。",
    "6. 设计并实现 Agent 工具调用链路追踪（基于 Micrometer Observation），观测工具调用耗时与成功率，支撑 Agent 效果评测。",
    "7. 参与会话、AI 意图等模块开发，保障多轮对话与意图路由稳定可用。",
]


def clear_runs(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)


def set_text(paragraph, text, bold=None):
    clear_runs(paragraph)
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold


def insert_after(paragraph, template, text):
    new_p = deepcopy(template._p)
    for child in list(new_p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph):
    p = paragraph._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def main():
    doc = Document(str(SRC))
    title_i = next(i for i, p in enumerate(doc.paragraphs) if "相闻" in p.text and "交通问答" in p.text)
    end_i = next(i for i, p in enumerate(doc.paragraphs) if i > title_i and "五育" in p.text)

    set_text(doc.paragraphs[title_i], TITLE, bold=True)
    set_text(doc.paragraphs[title_i + 1], OVERVIEW, bold=False)
    set_text(doc.paragraphs[title_i + 2], TECH, bold=False)

    paras = list(doc.paragraphs)
    for i in range(end_i - 1, title_i + 2, -1):
        delete_paragraph(paras[i])

    title_i = next(i for i, p in enumerate(doc.paragraphs) if "相闻" in p.text and "交通问答" in p.text)
    tech_p = doc.paragraphs[title_i + 2]
    item_tpl = tech_p
    for p in doc.paragraphs:
        if p.text.strip().startswith("1.规则引擎"):
            item_tpl = p
            break

    last = tech_p
    for text in DUTIES:
        last = insert_after(last, item_tpl, text)

    doc.save(str(SRC))
    doc.save(str(DL))
    print("DOCX saved")
    for p in doc.paragraphs:
        if any(k in p.text for k in ("相闻", "分片", "混合", "拒答", "性能", "提示词", "Micrometer", "会话", "召回")):
            if p.text.strip().startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.")) or "相闻" in p.text:
                print(">", p.text[:110])

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        for out in (OUT_PDF, OUT_PDF_MAIN):
            try:
                if out.exists():
                    out.unlink()
            except PermissionError:
                print("skip locked", out.name)
                continue
            d = word.Documents.Open(str(SRC), ReadOnly=True)
            d.SaveAs(str(out), FileFormat=17)
            d.Close(False)
            print("PDF", out.name, out.stat().st_size)
    finally:
        word.Quit()


if __name__ == "__main__":
    main()
