# -*- coding: utf-8 -*-
"""生成 BadCase Doctor 项目介绍 Word 文档（面试用 · 理念导向）。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)


def add_para(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.4


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)
        p.paragraph_format.space_after = Pt(5)


def add_highlight(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=(0x1A, 0x56, 0x8C))


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.6)
    s.right_margin = Cm(2.6)

    # 封面
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(t.add_run("BadCase Doctor"), size=28, bold=True, color=(0x1A, 0x56, 0x8C))
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub.add_run("AI 主动操作浏览器的智能测试平台"), size=15, color=(0x44, 0x44, 0x44))
    doc.add_page_break()

    # 一、我在做什么
    add_heading(doc, "一、我在做什么", 1)
    add_para(
        doc,
        "这不是一个传统的缺陷管理系统。"
        "我做的是：让 AI 像测试工程师一样，主动打开浏览器、按意图或用例去点、去填、去断言；"
        "测完之后，把发现的问题沉淀下来，再用统一的方式管理和继续处理。",
    )
    add_highlight(doc, "测，是核心；Bug、用例、BadCase，是测完之后的东西。")
    add_para(doc, "三者分别是什么：", bold=True)
    add_bullets(
        doc,
        [
            "测试用例（TestCase）：描述「要怎么测」——步骤、页面、预期。",
            "缺陷（Bug）：浏览器测下来发现的功能、交互类问题。",
            "BadCase：大模型类产品测下来发现的对话质量、准确率类问题。",
        ],
    )
    add_para(
        doc,
        "它们不是产品的出发点，而是测试执行的自然结果。"
        "产品要解决的是：谁来测、怎么测、测完怎么在一个地方说清楚、改清楚、追清楚。",
    )

    # 二、四条理念
    add_heading(doc, "二、产品理念", 1)
    add_para(doc, "下面四条是我设计整个产品的底层原则，功能都围绕它们长出来。", indent=0)

    ideas = [
        (
            "1. 所有东西都可以被对话操作",
            "用户不该记菜单、记表名、记按钮在哪。"
            "改负责人、搜缺陷、批量处理、发起一轮测试——都应该能说一句话就做完。"
            "对话不是聊天附件，而是产品的主操作方式。",
        ),
        (
            "2. 问题的所有关联业务，都应进入同一个上下文",
            "围绕一个问题，Agent 应该同时看到：相关缺陷、用例、BadCase、看板卡片、"
            "评论、日志、终端输出、历史记忆——而不是每次只盯着一张表。"
            "测试和质量本质是关联的，上下文也该是关联的。",
        ),
        (
            "3. Everything is Bash",
            "终端和浏览器一样，都是「执行面」。"
            "装环境、跑脚本、拉日志、调命令行工具，都走 bash；"
            "能写成命令的，就不硬造一套孤立 UI。"
            "降低自动化的门槛，也让 Agent 的能力边界更清晰。",
        ),
        (
            "4. 卡片（Card）抽象统一所有业务",
            "Bug、BadCase、用例在底层口径不同，但对人应该是一张「卡片」："
            "同一套迭代、状态、负责人、对话入口。"
            "用户跟卡片打交道，不必时刻想「我现在改的是哪张表」。"
            "卡片是统一抽象，底层各存各的，是工程实现。",
        ),
    ]
    for title, body in ideas:
        add_para(doc, title, bold=True)
        add_para(doc, body, indent=0.5)
        doc.add_paragraph()

    add_table(
        doc,
        ["理念", "一句话"],
        [
            ["对话操作一切", "说人话就能干活"],
            ["同一上下文", "一个问题，全貌可见"],
            ["Everything is Bash", "命令行与 GUI 同等重要"],
            ["Card 统一业务", "一种卡片，管所有类型"],
        ],
    )

    # 三、产品怎么运转
    add_heading(doc, "三、产品怎么运转", 1)
    add_para(doc, "从用户视角，主链路是这样的：", bold=True)
    add_bullets(
        doc,
        [
            "说意图或用例：「按 TC-001 测登录」「去测忘记密码流程」。",
            "AI 驱动浏览器：自动点击、输入、截图、比对预期。",
            "产出结果：候选缺陷、BadCase、执行记录，进入审核或入库。",
            "继续治理：用对话检索、修改、指派；用终端跑脚本深挖；用卡片在迭代里跟踪。",
        ],
    )
    add_highlight(doc, "浏览器负责「测」；对话 + 卡片 + 终端负责「管」和「继续干」。")

    add_heading(doc, "三个执行面", 2)
    add_table(
        doc,
        ["执行面", "做什么", "典型场景"],
        [
            ["浏览器", "模拟真人操作页面", "跑用例、复现步骤、发现 UI 问题"],
            ["对话（Agent）", "理解意图、编排动作", "搜缺陷、改状态、触发测试"],
            ["Bash（终端）", "命令行级执行与探查", "拉日志、跑脚本、调本地工具"],
        ],
    )
    add_para(
        doc,
        "三者不是三个产品，而是同一个 Agent 在不同面上的延伸。"
        "围绕一张卡片，可以一边对话改状态，一边开终端查日志，一边让浏览器再跑一遍复现。",
    )

    # 四、核心能力（不讲实现）
    add_heading(doc, "四、核心能力", 1)

    add_heading(doc, "4.1 浏览器主动测试", 2)
    add_bullets(
        doc,
        [
            "按测试用例自动执行步骤，像人一样操作页面。",
            "测完自动生成待确认的缺陷列表，人审核后再正式入库。",
            "也支持对话直接发起：「帮我把某某流程测一遍」。",
            "对大模型应用，可做对话准确率测试，产出 BadCase。",
        ],
    )

    add_heading(doc, "4.2 测试产物的统一管理", 2)
    add_bullets(
        doc,
        [
            "缺陷、BadCase、用例在同一项目、同一迭代计划下管理。",
            "状态流转、负责人、评论、通知——经典测试管理该有的都有。",
            "但入口统一在卡片和对话上，而不是三套割裂系统。",
        ],
    )

    add_heading(doc, "4.3 对话式治理", 2)
    add_bullets(
        doc,
        [
            "自然语言检索：「负责人是谁的 open 缺陷」「登录相关的问题」。",
            "自然语言修改：改标题、改状态、批量指派——改之前可预览、可采纳。",
            "多轮 Agent：能规划步骤、执行工具、总结观察，不是一问一答。",
        ],
    )

    add_heading(doc, "4.4 统一上下文与记忆", 2)
    add_bullets(
        doc,
        [
            "跨类型检索：一次搜索可命中缺陷、BadCase、用例、卡片。",
            "项目级记忆：打开项目时加载近期上下文，对话更连贯。",
            "为「一个问题看全貌」服务，而不是为单表查询服务。",
        ],
    )

    add_heading(doc, "4.5 桌面工作台", 2)
    add_bullets(
        doc,
        [
            "Electron 桌面端：项目、看板、Agent 对话、内嵌终端在一个工作台里。",
            "终端不是摆设，是与浏览器并列的日常操作区。",
            "适合测试工程师本机长时间使用，而不是偶尔打开的网页。",
        ],
    )

    # 五、规划方向
    add_heading(doc, "五、规划方向", 1)
    add_para(doc, "当前与接下来，按优先级大致是：", bold=True)
    add_bullets(
        doc,
        [
            "把浏览器主动测试做扎实：用例驱动、对话触发、审核入库闭环。",
            "强化「同一上下文」：围绕卡片一键装配缺陷、用例、日志、记忆。",
            "三执行面打通：同一张卡片上，对话、终端、浏览器由 Agent 自动选面。",
            "卡片进一步成为唯一心智模型，弱化用户对底层类型的感知。",
        ],
    )
    add_highlight(doc, "先能自动测 → 再管好测出来的 → 最后用对话、Bash、卡片拧成一根绳。")

    # 六、和常见工具的差异
    add_heading(doc, "六、和常见工具的差异", 1)
    add_table(
        doc,
        ["维度", "常见做法", "BadCase Doctor"],
        [
            ["起点", "先建缺陷库，自动化是附加", "先能测，缺陷是用例跑出来的"],
            ["操作方式", "表单 + 菜单", "对话为主，表单为辅"],
            ["实体关系", "各模块割裂", "同一上下文、Card 统一"],
            ["执行能力", "CI 脚本或手工", "浏览器 + Bash + Agent 三执行面"],
            ["AI 角色", "聊天问答", "真正参与测、找、改、追"],
        ],
    )

    # 结尾
    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        foot.add_run("本项目由本人独立设计与开发"),
        size=10,
        color=(0x88, 0x88, 0x88),
    )

    return doc


def main():
    docs = Path(__file__).resolve().parents[1] / "docs"
    docs.mkdir(parents=True, exist_ok=True)
    primary = docs / "BadCaseDoctor_项目介绍_面试用.docx"
    doc = build_doc()
    try:
        doc.save(str(primary))
        print(f"已生成: {primary}")
    except PermissionError:
        fallback = docs / "BadCaseDoctor_项目介绍_面试用_新版.docx"
        doc.save(str(fallback))
        print(f"原文件被占用，已另存为: {fallback}")


if __name__ == "__main__":
    main()
