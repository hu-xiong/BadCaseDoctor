from pathlib import Path

from docx import Document


path = Path(r"C:\Users\h2629\Downloads\胡雄java开发简历.docx")
doc = Document(str(path))

print("paragraphs", len(doc.paragraphs))
for i, para in enumerate(doc.paragraphs):
    text = " ".join(para.text.split())
    if text:
        print(f"P{i}: {text}")

print("tables", len(doc.tables))
for ti, table in enumerate(doc.tables):
    print(f"TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}")
    for ri, row in enumerate(table.rows):
        vals = [" ".join(cell.text.split()) for cell in row.cells]
        if any(vals):
            print(f"  R{ri}: {vals}")
