# -*- coding: utf-8 -*-
"""生成 BadCase Doctor 商业化项目说明书 Word 文档。"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "BadCaseDoctor_Business_Project_Brief.docx"

BLUE = RGBColor(26, 86, 140)
DARK = RGBColor(38, 38, 38)
GRAY = RGBColor(95, 95, 95)
LIGHT_BLUE = "D9EAF7"


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, size=9.5, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold, color=color or DARK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def title(doc, text, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_font(r, size=28, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run(subtitle)
    set_font(r, size=15, color=GRAY)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK)


def para(doc, text, bold=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(7)
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color or DARK)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, size=10.5, color=DARK)


def numbered(doc, items):
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{idx}. {item}")
        set_font(r, size=10.5, color=DARK)


def quote(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = table.cell(0, 0)
    shade(c, LIGHT_BLUE)
    cell_text(c, text, bold=True, size=11, color=BLUE)
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
    doc.add_paragraph()


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, LIGHT_BLUE)
        cell_text(c, h, bold=True, size=10, color=BLUE)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cell_text(cells[i], v)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.4)
    s.right_margin = Cm(2.4)

    title(doc, "BadCase Doctor 商业化项目说明书", "面向 AI 质量治理与智能测试的工作台产品")
    quote(doc, "一句话定位：BadCase Doctor 是帮助企业发现、管理、复盘并持续优化 AI 应用问题的质量治理平台。")
    para(doc, "适用对象：AI 应用团队、测试团队、研发团队、客服质检团队、企业数字化部门。", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, size=10)
    para(doc, "文档用途：项目介绍、融资沟通、合作洽谈、商业化方案说明。", align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, size=10)
    doc.add_page_break()

    heading(doc, "一、项目概述")
    para(doc, "BadCase Doctor 不是传统意义上的缺陷管理系统，也不是普通的 AI 聊天工具。它面向正在落地大模型应用的团队，解决 AI 产品在真实业务中“问题难发现、原因难定位、改进难闭环”的核心痛点。")
    para(doc, "项目将 BadCase、Bug、测试用例、执行记录、评论协作、AI 分析、终端执行和项目管理放在同一个工作台里，让团队可以围绕一个问题完成从发现、记录、分析、分派、修复、验证到复盘的全过程。")
    quote(doc, "核心价值：把 AI 应用质量问题从“靠人工感觉”变成“可发现、可追踪、可分析、可改进”的工程化流程。")

    heading(doc, "二、市场背景与行业机会")
    para(doc, "越来越多企业开始接入大模型，用于客服、办公助手、知识库问答、代码辅助、运营自动化、内部流程审批等场景。但 AI 应用上线后，企业很快会遇到一个共同问题：模型能回答，但回答是否稳定、是否准确、是否符合业务要求，很难长期保证。")
    bullets(doc, [
        "AI 输出具有不确定性：同一个问题在不同上下文、不同模型、不同时间可能出现不同结果。",
        "传统测试工具更擅长测固定流程，不擅长评估对话质量、推理质量和回答可信度。",
        "BadCase 往往散落在聊天记录、工单、截图、微信群、测试文档中，无法形成可复用资产。",
        "企业想优化模型效果，但缺少系统化数据，难以判断问题来自模型、提示词、知识库、工具调用还是业务逻辑。",
        "AI 产品迭代速度快，质量治理如果仍依赖人工表格和零散沟通，会拖慢上线节奏。",
    ])
    para(doc, "这意味着 AI 应用规模化之后，会自然产生一个新的基础能力需求：AI 质量治理。BadCase Doctor 的机会就在这里。")

    heading(doc, "三、要解决的问题")
    table(doc, ["问题", "企业现状", "BadCase Doctor 的解决方式"], [
        ["问题发现难", "靠人工体验、用户投诉、测试人员零散记录", "通过测试用例、对话触发、Agent 执行和 SDK 指标采集持续发现问题"],
        ["问题描述不清", "截图、文字、复现步骤分散，后续人员难理解", "统一沉淀为卡片，包含上下文、评论、执行结果和处理记录"],
        ["原因归因难", "不知道是模型、提示词、知识库还是工具链问题", "围绕 BadCase 采集模型、输入、外部依赖、服务稳定性、输出质量等因素"],
        ["协作效率低", "测试、研发、产品、算法之间靠人工同步", "项目、成员、负责人、状态流转、评论、通知形成闭环"],
        ["复盘资产少", "问题解决后没有沉淀，下次还会重复出现", "BadCase、Bug、TestCase 统一归档，变成可检索、可复用的质量资产"],
    ])

    heading(doc, "四、产品定位")
    para(doc, "BadCase Doctor 的产品定位可以概括为三个层次：")
    numbered(doc, [
        "AI 应用质量治理平台：用于记录、分析、追踪大模型应用中的 BadCase，帮助团队持续提升回答质量和业务稳定性。",
        "智能测试工作台：支持测试用例、缺陷、BadCase、浏览器测试、终端执行和 AI 对话式操作，提升测试与验证效率。",
        "面向未来的 AI 运维入口：通过 SDK、指标看板、日志和执行链路，把线上质量问题纳入长期监控和治理。",
    ])
    para(doc, "它不是只服务单个研发人员的小工具，而是可以面向团队、项目和企业流程落地的工作台产品。")

    heading(doc, "五、核心功能")
    heading(doc, "5.1 BadCase 全流程管理", 2)
    bullets(doc, ["记录 AI 对话或业务执行中不符合预期的案例。", "支持优先级、状态、负责人、评论、附件和复盘说明。", "支持从发现问题到修复验证的完整流转。", "把零散问题沉淀为企业自己的质量资产库。"])
    heading(doc, "5.2 Bug、BadCase、测试用例统一工作台", 2)
    bullets(doc, ["Bug 用于管理功能缺陷，BadCase 用于管理 AI 质量问题，测试用例用于定义验证标准。", "三类对象在同一项目、同一迭代、同一协作流程中管理。", "通过统一卡片视角降低使用门槛，让非技术人员也能理解问题进度。"])
    heading(doc, "5.3 AI 对话式操作", 2)
    bullets(doc, ["用户可以用自然语言查询、修改、创建、指派和总结问题。", "例如：查一下登录相关的未解决问题；把这个缺陷指派给张三；总结本周高频 BadCase。", "相比传统表单操作，对业务人员和测试人员更友好。"])
    heading(doc, "5.4 智能分析与原因归因", 2)
    bullets(doc, ["对 BadCase 进行原因分析，辅助判断问题来源。", "支持从模型参数、输入质量、知识库召回、工具调用、服务稳定性、输出质量等维度分析。", "帮助团队从发现一个问题升级到找到一类问题的根因。"])
    heading(doc, "5.5 SDK 与指标采集", 2)
    bullets(doc, ["计划支持 Java 和 Python SDK，让企业应用以较低成本接入质量观测。", "采集请求耗时、错误率、BadCase 类型、工具调用情况、流式输出体验等指标。", "与 Prometheus 等监控体系集成，为后续运维和看板分析提供数据基础。"])
    heading(doc, "5.6 桌面端与本地执行能力", 2)
    bullets(doc, ["提供桌面工作台，适合测试、研发、运维人员长期使用。", "集成终端和本地代理能力，便于拉日志、执行脚本、复现问题。", "后续可扩展浏览器自动测试、本地环境测试和远程 Agent 调度。"])

    heading(doc, "5.7 自动测试与浏览器执行", 2)
    para(doc, "自动测试是 BadCase Doctor 后续商业化能力中的关键模块。它不是简单录制脚本，而是让 AI Agent 根据测试用例或用户口头描述，像测试工程师一样打开浏览器、操作页面、观察结果、判断是否符合预期。")
    bullets(doc, [
        "测试用例驱动：用户维护结构化测试用例，例如登录、注册、支付、搜索、审批等流程，系统按步骤自动执行。",
        "口头描述驱动：用户直接说“帮我测一下登录失败提示是否正常”“去跑一遍忘记密码流程”，Agent 自动拆解步骤并执行。",
        "浏览器工具驱动：集成 browser-use、Playwright 或同类浏览器自动化能力，完成点击、输入、跳转、截图、等待、断言等操作。",
        "执行结果沉淀：测试失败时自动生成候选 Bug 或 BadCase，并附带页面截图、操作步骤、实际结果和预期差异。",
        "人工确认闭环：AI 发现的问题先进入预览或待确认状态，由测试或产品人员确认后正式入库，避免误报污染问题库。",
        "回归验证：当问题修复后，可自动重新执行关联测试用例，验证修复是否有效，减少人工回归成本。",
    ])
    quote(doc, "自动测试的目标：用户说清楚要测什么，AI 负责打开浏览器跑流程、找问题、留证据、进闭环。")

    heading(doc, "六、用户价值")
    table(doc, ["用户角色", "核心收益"], [
        ["企业管理者", "降低 AI 应用上线后的质量风险，让 AI 项目从试点走向可持续运营"],
        ["产品经理", "快速了解高频问题、用户痛点和版本质量，形成需求优化依据"],
        ["测试团队", "从手工记录问题转向用例化、流程化、智能化的质量验证"],
        ["研发团队", "获得更清晰的问题上下文、复现信息和优先级，减少沟通成本"],
        ["算法/提示词团队", "基于真实 BadCase 做模型、提示词、知识库和工具链优化"],
        ["运维团队", "通过指标、日志和告警关注 AI 应用稳定性，提前发现质量波动"],
    ])

    heading(doc, "七、商业价值与变现方向")
    para(doc, "BadCase Doctor 的商业化逻辑不是卖一个缺陷列表工具，而是成为企业 AI 应用长期运营中的质量基础设施。")
    heading(doc, "7.1 可收费模块", 2)
    bullets(doc, [
        "团队版订阅：项目管理、成员协作、BadCase 管理、测试用例管理、基础 AI 分析。",
        "企业版部署：私有化部署、权限体系、审计日志、组织级数据隔离、专属模型配置。",
        "质量观测套件：SDK、Prometheus 指标接入、质量看板、告警规则、趋势分析。",
        "智能测试套件：浏览器自动测试、测试集批量执行、回归验证、执行报告。",
        "行业模板服务：客服、电商、金融、政企知识库、代码助手等场景的 BadCase 分类模板。",
        "咨询与交付服务：帮助企业建立 AI 应用质量评估体系、BadCase 标注规范和运营流程。",
    ])
    heading(doc, "7.2 收费模式建议", 2)
    table(doc, ["模式", "适用客户", "说明"], [
        ["SaaS 订阅", "中小团队、AI 创业公司", "按成员数、项目数、执行次数、存储量收费"],
        ["私有化部署", "中大型企业、数据敏感行业", "一次性交付费 + 年度维护费"],
        ["用量计费", "测试执行量较大的客户", "按测试集执行次数、Agent 调用次数、指标采集量收费"],
        ["增值服务", "需要体系建设的客户", "模板定制、流程咨询、模型评测方案、运维对接"],
    ])

    heading(doc, "八、竞争差异")
    table(doc, ["对比对象", "常见能力", "BadCase Doctor 的差异"], [
        ["传统缺陷管理工具", "记录 Bug、分派、状态流转", "进一步覆盖 AI BadCase、质量归因、对话式操作和智能测试"],
        ["普通自动化测试平台", "执行固定脚本、生成报告", "更关注 AI 输出质量、对话质量和问题闭环，而不是只测页面流程"],
        ["大模型评测工具", "离线评测、打分、排行榜", "不仅评测，还把问题进入项目协作、修复、复测和运维流程"],
        ["聊天式 AI 助手", "回答问题、生成内容", "AI 是执行者和分析者，能围绕项目数据进行真实操作"],
    ])

    heading(doc, "九、当前阶段")
    para(doc, "当前项目已经具备较完整的产品雏形，不是停留在概念阶段。现有能力包括用户与项目管理、BadCase 管理、Bug 与测试用例相关模块、AI 对话与 Agent 工具调用、桌面端工作台、本地代理、沙箱预览、SDK 雏形和多项需求设计文档。")
    para(doc, "从商业化角度看，当前阶段适合定位为 MVP 到产品化早期阶段：核心方向清晰，功能边界已经展开，下一步重点应从功能可用转向场景打磨、稳定性、交付能力和标杆客户验证。")

    heading(doc, "十、二期规划")
    para(doc, "二期目标是从能管理 BadCase 升级为能持续发现 BadCase、解释 BadCase、推动 BadCase 解决。")
    numbered(doc, [
        "智能测试闭环：建设测试集、测试用例批量执行、自动评分、失败自动生成 BadCase 的完整链路。",
        "浏览器执行能力：支持 Agent 借助 browser-use、Playwright 等工具在浏览器中执行操作，自动复现流程、截图、采集证据。",
        "口头测试能力：支持用户用自然语言描述测试目标，系统自动拆解为可执行步骤，并将结果沉淀为测试报告。",
        "质量看板：按项目、版本、模型、场景统计 BadCase 数量、类型、修复周期和复发率。",
        "SDK 正式化：完善 Java/Python SDK，降低企业应用接入门槛。",
        "Prometheus 集成：将请求耗时、错误率、流式输出、工具调用、BadCase 类型等指标接入监控体系。",
        "知识沉淀：形成 BadCase 分类、归因标签、处理经验和行业模板。",
        "权限与审计：完善企业级成员权限、操作审计、数据隔离和安全控制。",
    ])

    heading(doc, "十一、运维规划")
    para(doc, "项目商业化后，运维能力会直接影响客户信任。建议将运维体系分为产品运维、系统运维和质量运维三层。")
    heading(doc, "11.1 产品运维", 2)
    bullets(doc, ["建立客户项目空间、成员权限、版本发布、使用统计和问题反馈通道。", "提供客户成功流程：首次接入、模板配置、项目导入、成员培训、月度质量报告。", "围绕 BadCase 数量、解决周期、复发率等指标证明产品价值。"])
    heading(doc, "11.2 系统运维", 2)
    bullets(doc, ["建设服务健康监控：接口可用性、响应耗时、错误率、任务队列、数据库状态。", "建设数据备份机制：项目数据、BadCase 数据、评论附件、测试报告定期备份。", "建设日志与审计：关键操作留痕，支持问题追踪和企业合规要求。", "支持 SaaS 与私有化两套部署路径，满足不同客户的数据安全诉求。"])
    heading(doc, "11.3 AI 质量运维", 2)
    bullets(doc, ["持续跟踪模型效果波动，识别某个模型、某类问题或某个业务场景的质量下降。", "对高频 BadCase 自动聚类，输出优化建议。", "建立提示词、知识库、工具链变更后的回归测试机制。", "将 AI 质量从一次性评测变成日常运营指标。"])

    heading(doc, "十二、远期规划：Coding 与智能研发")
    para(doc, "远期方向可以从 AI 质量治理延伸到智能研发与智能运维，形成发现问题、定位问题、生成修复建议、验证修复结果的闭环。")
    bullets(doc, [
        "从描述到测试：用户只需要描述业务目标，Agent 自动生成测试步骤、驱动浏览器执行，并沉淀为可复用测试用例。",
        "智能 Coding 助手：围绕具体 Bug 或 BadCase，自动读取上下文、定位代码、生成修复建议。",
        "代码变更预览：所有 AI 生成修改先进入差异预览，由人确认后采纳。",
        "自动回归验证：修复后自动运行相关测试用例，确认问题是否解决。",
        "研发知识库：沉淀项目结构、历史修复记录、常见问题和最佳实践。",
        "从测试到修复闭环：不只告诉团队哪里坏了，还帮助团队推进怎么修、修完怎么验。",
        "多 Agent 协作：测试 Agent、分析 Agent、Coding Agent、运维 Agent 分工协作，形成企业内部 AI 研发助手体系。",
    ])
    quote(doc, "远期愿景：成为企业 AI 应用和软件研发流程中的质量中枢，而不是一个孤立的问题记录工具。")

    heading(doc, "十三、里程碑建议")
    table(doc, ["阶段", "目标", "关键交付"], [
        ["MVP 打磨", "验证核心场景", "BadCase/Bug/用例管理、AI 对话操作、基础项目协作"],
        ["二期产品化", "形成质量闭环", "测试集、自动评分、浏览器自动测试、口头测试、质量看板、SDK、Prometheus 接入"],
        ["商业试点", "服务真实客户", "私有化部署、客户模板、月度质量报告、客户成功流程"],
        ["规模化", "成为质量基础设施", "多租户、权限审计、行业方案、智能运维、Coding Agent"],
    ])

    heading(doc, "十四、投资人视角总结")
    para(doc, "BadCase Doctor 的价值在于抓住了 AI 应用规模化之后必然出现的质量治理需求。随着企业越来越多地把大模型接入真实业务，质量问题会从偶尔出现的体验问题变成影响业务可信度、客户满意度和运营成本的系统问题。")
    para(doc, "传统缺陷管理工具无法很好覆盖 AI 输出质量，普通评测工具又难以进入日常项目协作和修复流程。BadCase Doctor 选择的切入点是把 BadCase 管理、智能测试、质量归因、指标观测和后续修复协作放在一个工作台中，形成从发现到解决的闭环。")
    para(doc, "如果二期能够把 SDK、指标看板、测试集自动执行和企业级部署能力做稳定，就具备向 SaaS 订阅、私有化交付和行业解决方案延展的商业化基础。")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("BadCase Doctor：让 AI 质量问题可管理、可解释、可运营、可持续改进。")
    set_font(r, size=12, bold=True, color=BLUE)
    return doc


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    try:
        doc.save(str(OUT))
        print(OUT)
    except PermissionError:
        fallback = OUT.with_name("BadCaseDoctor_Business_Project_Brief_fixed.docx")
        doc.save(str(fallback))
        print(fallback)


if __name__ == "__main__":
    main()
