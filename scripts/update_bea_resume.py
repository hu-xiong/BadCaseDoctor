# -*- coding: utf-8 -*-
"""完善东亚银行项目经历，并导出 PDF。"""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from copy import deepcopy

SRC = Path(r"C:\Users\h2629\Downloads\胡雄java开发简历.docx")
OUT_DOCX = Path(r"C:\Users\h2629\Desktop\胡雄java开发简历_东亚银行完善.docx")
OUT_PDF = Path(r"C:\Users\h2629\Desktop\胡雄java开发简历_东亚银行完善.pdf")

# 替换原 118~121，并追加职责条目
NEW_AFTER_TITLE = [
    "东亚银行手机银行与官网支撑系统，覆盖转账汇款、支付、金融理财、会员管理、财务审批与资金流转等核心业务。含会员积分与「万里行」类权益体系（消费/活动累计积分、积分兑换、权益发放与核销），支持香港地区青年派发港币等活动能力，并落地 App 端二维码电子支付，保障 Web 与 App 业务协同运转。",
    "技术栈: Spring Boot, Oracle, Ionic, Angular, RESTful API",
    "负责的工作：",
    "1.负责 Web 端业务模块日常开发与维护，跟进转账、支付、会员、理财等接口与页面联调，处理线上问题与需求迭代。",
    "2.参与会员积分 / 「万里行」类权益模块：积分规则配置、消费/活动累计、积分查询与流水、兑换与核销、权益到账状态同步，保证积分增减可追溯。",
    "3.参与 App 手机银行香港地区二维码电子支付：支付下单、扫码/被扫、结果回调与订单状态同步，保障支付闭环。",
    "4.参与财务审批与支付流转：按审批节点推进单据状态，保证审批流与资金操作可审计。",
    "5.基于 Oracle 完成积分流水、订单、会员等数据的查询与持久化，处理事务一致性与常见慢查询。",
    "6.配合测试与业务方联调验收，沉淀接口与问题记录，保障版本按期上线。",
]


def set_run_font(run, bold=False):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.bold = bold
    if run.font.size is None:
        run.font.size = Pt(10.5)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)


def set_paragraph_text(paragraph, text, bold=False):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    return paragraph


def insert_paragraph_after(paragraph, text, bold=False):
    """在 paragraph 后插入同级段落。"""
    new_p = deepcopy(paragraph._p)
    # 清空 runs
    for child in list(new_p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    # wrap as paragraph
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    set_run_font(run, bold=bold)
    return new_para


def update_docx():
    doc = Document(str(SRC))
    cell = doc.tables[0].cell(2, 0)
    paras = cell.paragraphs

    title_idx = None
    for i, p in enumerate(paras):
        if "2018.7-2020.6" in p.text and "东亚银行" in p.text:
            title_idx = i
            break
    if title_idx is None:
        raise RuntimeError("未找到东亚银行项目标题段落")

    # 原内容：title + 118,119,120,121（共 4 段概述/技术/职责）
    # 先改写紧随其后的现有段落，不够再插入，多余清空
    existing_follow = []
    for j in range(title_idx + 1, len(paras)):
        # 只处理到文档末尾（东亚银行是最后一项）
        existing_follow.append(paras[j])

    # 保留标题加粗
    set_paragraph_text(paras[title_idx], "2018.7-2020.6          东亚银行手机银行", bold=True)

    # 用现有 follow 段落填充；不够则在最后一个后插入
    last = paras[title_idx]
    used = []
    for idx, text in enumerate(NEW_AFTER_TITLE):
        if idx < len(existing_follow):
            p = existing_follow[idx]
            set_paragraph_text(p, text, bold=False)
            last = p
            used.append(p)
        else:
            last = insert_paragraph_after(last, text, bold=False)
            used.append(last)

    # 清空多余旧段落
    for p in existing_follow[len(NEW_AFTER_TITLE) :]:
        clear_paragraph(p)

    doc.save(str(OUT_DOCX))
    print("DOCX:", OUT_DOCX)
    return OUT_DOCX


def export_pdf(docx_path: Path, pdf_path: Path):
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path))
        # 17 = wdFormatPDF
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close(False)
    finally:
        word.Quit()
    print("PDF:", pdf_path)


def main():
    docx_path = update_docx()
    try:
        export_pdf(docx_path, OUT_PDF)
    except Exception as e:
        print("Word 导出 PDF 失败，尝试备用方案:", e)
        # fallback: pip install 后用 docx2pdf 或提示用户
        try:
            from docx2pdf import convert

            convert(str(docx_path), str(OUT_PDF))
            print("PDF:", OUT_PDF)
        except Exception as e2:
            print("备用导出也失败:", e2)
            raise


if __name__ == "__main__":
    main()
